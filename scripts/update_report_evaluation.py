# -*- coding: utf-8 -*-
"""Обновление раздела оценки и пометок для картинок в оригинальном Word-отчёте."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

TARGET = Path(r"c:\Users\User\Downloads\teammate_finder_cherkashina_ktbo3_3.docx")
PROJECT = Path(r"C:\Users\User\Desktop\temmate_finder_project")

# Точные пути к готовым графикам
GRAPHS = {
    "ris6": PROJECT / "results" / "custom_model" / "precision_vs_threshold.png",
    "ris7": PROJECT / "results" / "custom_model" / "all_methods_comparison.png",
    "ris7_alt": PROJECT / "methods_comparison.png",
    "ris8": PROJECT / "results" / "custom_model" / "predicted_vs_actual.png",
    "ris8_alt": PROJECT / "results" / "custom_model" / "metrics_summary.png",
    "ris8_tsne": PROJECT / "tsne_comparison.png",
}

FIGURE_MARKERS = {
    "Рисунок 1 – Архитектура модели Transformer": (
        "▶ ВСТАВИТЬ ИЗОБРАЖЕНИЕ (Рис. 1): Схема Transformer (Encoder/Decoder, self-attention). "
        "Нарисовать в draw.io / PowerPoint или взять из учебных материалов. Теория §1.3."
    ),
    "Рисунок 2 – Размеченный датасет": (
        "▶ ВСТАВИТЬ СКРИНШОТ (Рис. 2): Файл data/gaming_compatibility_dataset.csv в Excel или VS Code. "
        "Видны столбцы description1, description2, compatibility_score, game_type."
    ),
    "Рисунок 3 – метод очистки": (
        "▶ ВСТАВИТЬ СКРИНШОТ (Рис. 3): preprocessing/text_processor.py — метод clean_text() в IDE."
    ),
    "Рисунок 4 – процесс замены сленга": (
        "▶ ВСТАВИТЬ ИЗОБРАЖЕНИЕ (Рис. 4): Таблица «до/после» сленга "
        "(«агр сап» → «агрессивный саппорт») или скрин gaming_slang_dict."
    ),
    "Рисунок 5 – Тестовая база": (
        "▶ ВСТАВИТЬ СКРИНШОТ (Рис. 5): teammate_finder.db в DB Browser — таблица users "
        "(после python scripts/seed_database.py)."
    ),
    "Рисунок 6 – Зависимость точности": (
        f"▶ ВСТАВИТЬ ГРАФИК (Рис. 6): {GRAPHS['ris6']}. "
        "График Precision vs порог для СОБСТВЕННОЙ модели на ОТЛОЖЕННОМ test.csv (11 пар). "
        "Сгенерировать: python scripts/evaluate_custom_model.py"
    ),
    "Рисунок 7 – сравнение моделей": (
        f"▶ ВСТАВИТЬ ГРАФИК (Рис. 7): {GRAPHS['ris7']} "
        f"(или {GRAPHS['ris7_alt']} — только baseline TF-IDF / SBERT / mBERT). "
        "Столбчатая диаграмма Precision и NDCG."
    ),
    "Рисунок 8 – корреляция": (
        f"▶ ВСТАВИТЬ ГРАФИК (Рис. 8): {GRAPHS['ris8']} — scatter «предсказание vs экспертная оценка» "
        f"на test.csv. Дополнительно (по желанию): {GRAPHS['ris8_tsne']} для baseline-методов."
    ),
    "Рисунок 9 – Пользовательский запрос": (
        "▶ ВСТАВИТЬ СКРИНШОТ (Рис. 9): Swagger http://localhost:8000/docs — POST /similarity/search "
        "или форма поиска в teammate_finder.html."
    ),
    "Рисунок 10 – Препроцессинг": (
        "▶ ВСТАВИТЬ СКРИНШОТ (Рис. 10): Пример препроцессинга — текст со сленгом и результат clean_text."
    ),
    "Рисунок 11 – Поиск схожих": (
        "▶ ВСТАВИТЬ СКРИНШОТ (Рис. 11): database/sqlite_db.py — метод find_similar_users."
    ),
    "Рисунок 12 – Список наиболее схожих": (
        "▶ ВСТАВИТЬ СКРИНШОТ (Рис. 12): JSON-ответ /similarity/search в Swagger после Execute."
    ),
    "Рисунок 13 – Визуализация результатов": (
        "▶ ВСТАВИТЬ СКРИНШОТ (Рис. 13): teammate_finder.html — карточки тиммейтов с % совместимости."
    ),
}

TRAINING_MARKER = (
    "▶ ВСТАВИТЬ СКРИНШОТ (§3.2): Консоль обучения: python -m model.train --dataset data/train.csv --epochs 30. "
    "Видны строки Epoch 30/30, train/val loss, «сохранена лучшая модель»."
)

# Замена по уникальному началу абзаца
TEXT_REPLACEMENTS = {
    "Обучение выполнялось методом Triplet Loss (margin = 0.25) на размеченном датасете gaming_compatibility_dataset.csv": (
        "Обучение выполнялось методом Triplet Loss (margin = 0.25). Полный датасет gaming_compatibility_dataset.csv "
        "(53 пары) разделён на обучающую выборку train.csv (42 пары) и отложенную тестовую test.csv (11 пар) "
        "скриптом scripts/split_dataset.py. Модель обучалась только на train.csv. Из пар с compatibility_score ≥ 0.7 "
        "формировались триплеты (anchor, positive, negative). Обучение — 30 эпох; лучшее validation loss на валидационной "
        "части train — 0.0000 (эпоха 27). Параметров модели: ~1.58 млн, размер словаря: 206 слов."
    ),
    "Основу датасета составили более 50 пар описаний игровых стилей": (
        "Основу датасета составили 53 пары описаний игровых стилей, включая исходные 30 экспертно размеченных пар "
        "и дополнительные примеры с геймерским сленгом. Для честной оценки данные разделены: 42 пары — обучение "
        "(data/train.csv), 11 пар — отложенный тест (data/test.csv), не виденный модели при обучении. "
        "Критерии оценки включали ролевую синергию, стилевое соответствие и тактическую сочетаемость (рис. 5). "
        "Пары с оценкой ≥ 0.7 — положительные примеры Triplet Loss, пары с оценкой ≤ 0.45 — отрицательные."
    ),
    "Помимо обучения модели, датасет применялся для настройки порога косинусной схожести": (
        "Помимо обучения, на отложенной тестовой выборке (test.csv) подбирался порог косинусной схожести. "
        "Для API выбрано значение 0.35–0.40; для метрики Precision@0.5 (сопоставимо с baseline) "
        "использовался порог предсказания 0.5 (рис. 6)."
    ),
    "Датасет сбалансирован по игровым дисциплинам (Dota 2, CS:GO, Valorant) и ролевым амплуа": (
        "Датасет сбалансирован по игровым дисциплинам (Dota 2, CS:GO, Valorant) и ролевым амплуа "
        "(керри, саппорт, мидер). Собственная модель после обучения на train.csv корректно обрабатывает "
        "как литературные описания, так и разговорный сленг («агр сап», «каррю на лейте»)."
    ),
    "Оценка эффективности проводилась в два этапа. На первом этапе сравнивались готовые методы": (
        "Оценка эффективности проводилась в три этапа. "
        "Этап 1 — сравнение готовых методов (TF-IDF, Sentence-BERT, Multilingual-BERT) на 30 аннотированных парах. "
        "Этап 2 — обучение собственной модели TextEmbeddingModel только на train.csv (42 пары). "
        "Этап 3 — честная оценка на отложенной test.csv (11 пар), не участвовавших в обучении."
    ),
    "Сравнительный анализ готовых моделей показал precision 96.67% у Sentence-BERT": (
        "Сравнительный анализ готовых моделей показал Precision 96.67% у Sentence-BERT при пороге 0.5 "
        "и NDCG 0.9731 (рис. 7). TF-IDF показал Precision 63.33%, Multilingual-BERT — 90.00%. "
        "Для узкоспециализированной задачи с игровым сленгом принято решение о собственной архитектуре, "
        "обучаемой на размеченных парах совместимости."
    ),
    "NDCG 0.9731 свидетельствует о точном ранжировании результатов": (
        "На отложенной тестовой выборке (11 пар, test.csv) собственная модель TextEmbeddingModel показала: "
        "Precision 85.71% (порог pred≥0.5, actual≥0.7), NDCG 0.9713, корреляция Пирсона с экспертными оценками 0.2694, "
        "MAE 0.388, RMSE 0.521 (рис. 6, 8). Метрики получены скриптом scripts/evaluate_custom_model.py — "
        "модель не видела тестовые пары при обучении. NDCG 0.97 свидетельствует о точном ранжировании: "
        "наиболее совместимые пары получают более высокую предсказанную схожесть."
    ),
    "Практический вывод: для итоговой системы реализована собственная модель Embedding + BiLSTM + Attention": (
        "Практический вывод: реализована собственная модель Embedding + BiLSTM + Attention "
        "(128 измерений, ~1.58 млн параметров, словарь 206 слов), обученная Triplet Loss на train.csv. "
        "На отложенном тесте: Precision 85.71%, NDCG 0.9713. Модель компактнее трансформеров и учитывает "
        "игровой сленг через препроцессор и предметно-ориентированное обучение."
    ),
    "Дополнительно на тех же данных проводился сравнительный анализ baseline-методов": (
        "Дополнительно проводился сравнительный анализ baseline-методов (TF-IDF, Sentence-BERT, Multilingual-BERT) "
        "и собственной модели на одном графике (рис. 7, файл results/custom_model/all_methods_comparison.png). "
        "Это позволило сопоставить готовые трансформеры с обученной LSTM+Attention моделью."
    ),
}


def set_text(paragraph, text: str, highlight: bool = False):
    for hyperlink in list(paragraph._p.findall(qn("w:hyperlink"))):
        link_text = "".join(node.text or "" for node in hyperlink.iter(qn("w:t")))
        if re.fullmatch(r"приложение [1-5]", link_text.strip(), re.I):
            paragraph._p.remove(hyperlink)
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        run.bold = True


def insert_before(paragraph, text: str, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    np = Paragraph(new_p, paragraph._parent)
    if style is not None:
        np.style = style
    set_text(np, text, highlight=True)
    return np


def upsert_marker_before(paragraph, marker: str, style=None):
    prev = paragraph._p.getprevious()
    if prev is not None:
        prev_para = Paragraph(prev, paragraph._parent)
        if "▶ ВСТАВИТЬ" in prev_para.text:
            set_text(prev_para, marker, highlight=True)
            return
    insert_before(paragraph, marker, style)


def main():
    doc = Document(str(TARGET))
    body_style = None
    for p in doc.paragraphs:
        if p.text.startswith("Представьте") or p.text.startswith("В современном"):
            body_style = p.style
            break

    # Текстовые замены
    for p in doc.paragraphs:
        t = p.text
        for start, new_text in TEXT_REPLACEMENTS.items():
            if start in t:
                set_text(p, new_text)
                break

    # Пометки перед рисунками
    for p in doc.paragraphs:
        t = p.text.strip()
        for key, marker in FIGURE_MARKERS.items():
            if key in t:
                upsert_marker_before(p, marker, body_style)
                break

    # Пометка консоли обучения
    for p in doc.paragraphs:
        if p.text.strip().startswith("Оценка эффективности модели"):
            prev = p._p.getprevious()
            if prev is not None:
                prev_para = Paragraph(prev, p._parent)
                if "▶ ВСТАВИТЬ" in prev_para.text and "model.train" in prev_para.text:
                    set_text(prev_para, TRAINING_MARKER, highlight=True)
            break

    # Пометка перед Рис. 10 если отсутствует
    for p in doc.paragraphs:
        if "Рисунок 10" in p.text:
            upsert_marker_before(p, FIGURE_MARKERS["Рисунок 10 – Препроцессинг"], body_style)
            break

    doc.save(str(TARGET))
    print(f"Обновлён: {TARGET}")
    print("\nКуда вставлять картинки:")
    print("  Рис. 1  — схема Transformer (нарисовать)")
    print(f"  Рис. 2  — скрин data/gaming_compatibility_dataset.csv")
    print("  Рис. 3  — скрин text_processor.py")
    print("  Рис. 4  — таблица сленга до/после")
    print("  Рис. 5  — скрин SQLite users")
    print(f"  Рис. 6  — {GRAPHS['ris6']}")
    print(f"  Рис. 7  — {GRAPHS['ris7']}")
    print(f"  Рис. 8  — {GRAPHS['ris8']}")
    print("  Рис. 9  — Swagger или UI")
    print("  Рис. 10 — пример препроцессинга")
    print("  Рис. 11 — sqlite_db.py")
    print("  Рис. 12 — JSON ответ API")
    print("  Рис. 13 — teammate_finder.html")


if __name__ == "__main__":
    main()
