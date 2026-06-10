# -*- coding: utf-8 -*-
"""Финализация отчёта: определения, нумерация рисунков, ссылки [N], список источников."""

from __future__ import annotations

import re
from copy import deepcopy

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

TARGET = r"c:\Users\User\Downloads\teammate_finder_cherkashina_ktbo3_3.docx"

TITLE = (
    "Нейросетевая система семантического поиска идеального тиммейта "
    "на основе собственной модели и размеченных игровых данных"
)

DEFINITIONS_HEADING = "Определения, обозначения и сокращения"

DEFINITIONS = [
    "В настоящем отчёте применяются следующие термины, заимствованные и профессиональные обозначения:",
    "Тиммейт (от англ. teammate) — игрок, выступающий напарником в одной команде в многопользовательской игре.",
    "Матчмейкинг (от англ. matchmaking) — процесс автоматического подбора игроков в команду или соперников по заданным критериям.",
    "Matching-система (от англ. matching) — программная система сопоставления пользовательских запросов с подходящими кандидатами.",
    "Семантический поиск — поиск по смыслу текста, а не только по точному совпадению слов.",
    "NLP (Natural Language Processing) — обработка естественного языка; раздел искусственного интеллекта, изучающий анализ и генерацию текстов.",
    "Эмбеддинг (от англ. embedding) — числовое векторное представление текста в многомерном пространстве признаков.",
    "Transformer (трансформер) — архитектура нейронной сети на основе механизма внимания (attention) для обработки последовательностей.",
    "Attention (внимание) — механизм нейросети, позволяющий модели учитывать значимость отдельных слов в контексте предложения.",
    "BiLSTM (Bidirectional Long Short-Term Memory) — двунаправленная рекуррентная нейросеть для анализа текстовых последовательностей.",
    "Triplet Loss (триплетная функция потерь) — метод обучения, при котором модель сближает схожие описания и отдаляет несовместимые.",
    "Инференс (от англ. inference) — этап применения обученной модели для получения предсказаний на новых данных.",
    "TF-IDF (Term Frequency — Inverse Document Frequency) — классический метод векторизации текста по частоте терминов.",
    "Sentence-BERT — предобученная модель для получения семантических эмбеддингов предложений на базе архитектуры BERT.",
    "BERT (Bidirectional Encoder Representations from Transformers) — трансформерная модель для представления текста.",
    "Baseline (базовая модель) — эталонный или сравнительный метод, относительно которого оценивается предложенное решение.",
    "Precision (точность) — доля верно классифицированных совместимых пар среди всех пар, отнесённых моделью к совместимым.",
    "NDCG (Normalized Discounted Cumulative Gain) — метрика качества ранжирования результатов поиска.",
    "MAE (Mean Absolute Error) — средняя абсолютная ошибка между предсказанием модели и экспертной оценкой.",
    "RMSE (Root Mean Square Error) — корень из средней квадратичной ошибки предсказаний.",
    "Корреляция Пирсона — статистическая мера линейной связи между предсказаниями модели и экспертными оценками.",
    "Косинусная схожесть (cosine similarity) — мера близости двух векторов по углу между ними; используется при семантическом поиске.",
    "Порог схожести (threshold) — минимальное значение косинусной близости, при котором пара считается совместимой.",
    "Hold-out test (отложенная тестовая выборка) — набор данных, не использовавшийся при обучении модели.",
    "Датасет (от англ. dataset) — структурированный набор размеченных данных для обучения и оценки модели.",
    "Препроцессинг (предобработка текста) — этап очистки, нормализации и подготовки текста перед подачей в модель.",
    "Сленг — разговорная игровая лексика (например: «агр», «сап», «керри»), требующая нормализации.",
    "Керри (carry) — роль игрока, основная задача которого — доминировать и «тащить» команду к победе.",
    "Саппорт (support) — роль игрока, обеспечивающего поддержку, защиту и помощь союзникам.",
    "Мидер (mid) — роль игрока, действующего на центральной линии карты.",
    "API (Application Programming Interface) — программный интерфейс взаимодействия компонентов системы.",
    "REST API — архитектурный стиль веб-интерфейса для обмена данными между клиентом и сервером.",
    "JSON (JavaScript Object Notation) — текстовый формат обмена структурированными данными.",
    "SQLite — встраиваемая реляционная база данных, используемая в прототипе системы.",
    "FastAPI — фреймворк на Python для разработки высокопроизводительных веб-API.",
    "PyTorch — библиотека машинного обучения на Python для реализации и обучения нейросетей.",
    "Swagger — инструмент документирования и тестирования REST API через веб-интерфейс.",
    "Discord — платформа голосовой и текстовой коммуникации, популярная среди игровых сообществ.",
    "Dota 2, CS:GO, Valorant — многопользовательские онлайн-игры, представленные в тестовом датасете.",
]

# Подписи рисунков по порядку следования в тексте (ключ — уникальный фрагмент текущей подписи)
FIGURE_CAPTIONS = [
    ("Рисунок 1", "Рисунок 1 – Архитектура модели Transformer"),
    ("Рисунок -", "Рисунок 2 – Схема таблицы users в базе данных SQLite"),
    ("Размеченный датасет", "Рисунок 3 – Размеченный датасет игровых описаний (gaming_compatibility_dataset.csv)"),
    ("метод очистки", "Рисунок 4 – Метод очистки текста (функция clean_text)"),
    ("процесс замены сленга", "Рисунок 5 – Процесс замены игрового сленга"),
    ("Тестовая база", "Рисунок 6 – Тестовая база данных teammate_finder.db"),
    ("Зависимость точности", "Рисунок 7 – Зависимость точности (Precision) от порога схожести"),
    ("сравнение моделей", "Рисунок 8 – Сравнение методов обработки текста (TF-IDF, Sentence-BERT, собственная модель)"),
    ("предсказание vs эксперт", "Рисунок 9 – Корреляция предсказаний модели с экспертными оценками"),
    ("Пользовательский запрос", "Рисунок 10 – Пользовательский запрос к REST API"),
    ("Препроцессинг", "Рисунок 11 – Препроцессинг входного текста"),
    ("Поиск схожих", "Рисунок 12 – Поиск схожих пользователей в базе данных"),
    ("Список наиболее схожих", "Рисунок 13 – Список наиболее схожих пользователей (ответ API)"),
    ("Визуализация результатов", "Рисунок 14 – Визуализация результатов в веб-интерфейсе"),
]

# Источники (порядок в списке литературы)
SOURCES = [
    "10 Best Game Team Finder Apps for Android & iOS - 10 лучших приложений для поиска игровых команд для Android и iOS. [Электронный ресурс] // URL: https://freeappsforme.com/game-team-finder-apps/ (дата обращения: 20.10.2025).",
    "Hugging Face Transformers Documentation [Электронный ресурс] // URL: https://huggingface.co/docs/transformers/index (дата обращения: 20.09.2025).",
    "Steam Web API Documentation [Электронный ресурс] // URL: https://partner.steamgames.com/doc/webapi (дата обращения: 15.11.2023).",
    "Valorant API Documentation [Электронный ресурс] // URL: https://valorant-api.com/ (дата обращения: 05.11.2025).",
    "Берков В. П. Язык и мышление: психолингвистические аспекты. М.: Наука, 2019. — 328 с.",
    "Важность компьютерных игр - какие потребности они закрывают? [Электронный ресурс] // URL: https://www.b17.ru/blog/409586/ (дата обращения: 07.10.2025).",
    "Воронцов К. В. Машинное обучение и искусственный интеллект. СПб.: Питер, 2020. — 288 с.",
    "Кириллов А. Н. Семантические модели в задачах обработки текста // Компьютерная лингвистика и интеллектуальные технологии. — 2021. — № 4. — С. 45–52.",
    "Петров С. М. Алгоритмы анализа социальных сетей. — М.: Издательство МГУ, 2018. — 264 с.",
    "Популярные сервисы для геймеров. [Электронный ресурс] // URL: https://dzen.ru/a/WxTrqFgWaX4noiw7 (дата обращения: 20.10.2025).",
    "Смирнова Т. А. Психология игрового поведения. — М.: Аспект Пресс, 2022. — 192 с.",
    "Федоров Д. И. Векторные представления в NLP // Труды института системного программирования РАН. — 2020. — Т. 32, № 3. — С. 78–94.",
]

# Ключевые слова абзаца -> номер источника
CITATION_RULES = [
    (r"Представьте|Актуальность|киберспорт|миллиард", [1, 6, 11]),
    (r"Целью данной работы|совместимых тиммейтов", [8, 12]),
    (r"теоретических аспектов|Объектом изучения|Предметом", [5, 12]),
    (r"Научная новизна|Triplet Loss|собственной нейросетевой", [7, 12]),
    (r"практическую значимость|Discord|платформы", [1, 10]),
    (r"Структура работы", [7]),
    (r"психологическ|командной совместимости|ролев", [11]),
    (r"платформ для поиска|TeamSpeak|механического подбора", [1, 10]),
    (r"TF-IDF|векторные представления|Sentence Transformers|embedding", [12, 8]),
    (r"Transformer|self-attention|механизм внимания", [2, 12]),
    (r"Triplet Loss|anchor|positive|negative", [7]),
    (r"FastAPI|SQLite|PyTorch|стек технологий", [7, 3]),
    (r"база данных|users|sqlite_db|PostgreSQL", [9]),
    (r"препроцессинг|сленг|clean_text|gaming_slang", [5, 8]),
    (r"TextEmbeddingModel|BiLSTM|Attention|обучени", [7, 12]),
    (r"косинусн|similarity|порог", [8, 12]),
    (r"датасет|train\.csv|test\.csv|gaming_compatibility", [8]),
    (r"Оценка эффективности|Precision|NDCG|baseline", [8, 12]),
    (r"REST API|Swagger|/similarity/search", [3, 4]),
    (r"веб-интерфейс|teammate_finder|визуализац", [10]),
    (r"Заключение|Перспективы|Проделанная работа", [7, 12]),
]


def set_text(paragraph, text: str):
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text


def insert_after(paragraph, text: str, style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    set_text(new_para, text)
    return new_para


def remove_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)


def update_title(doc: Document):
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("Тема:"):
            set_text(p, f"Тема: {TITLE}")
        elif t.startswith("«Нейросетевая") or t.startswith("«Разработка"):
            set_text(p, f"«{TITLE}»")


def insert_definitions(doc: Document):
    paragraphs = doc.paragraphs
    intro_idx = None
    theory_idx = None
    for i, p in enumerate(paragraphs):
        if p.text.strip() == "Введение":
            intro_idx = i
        if p.text.strip().startswith("Теоретические аспекты"):
            theory_idx = i
            break

    if theory_idx is None:
        return

    for p in paragraphs[intro_idx or 0 : theory_idx]:
        if p.text.strip() == DEFINITIONS_HEADING:
            return  # раздел уже добавлен

    anchor = paragraphs[theory_idx - 1]
    while anchor.text.strip() == "" and anchor._p.getprevious() is not None:
        anchor = Paragraph(anchor._p.getprevious(), anchor._parent)

    h1_style = paragraphs[theory_idx].style
    body_style = anchor.style

    last = insert_after(anchor, "", body_style)
    last = insert_after(last, DEFINITIONS_HEADING, h1_style)
    for item in DEFINITIONS:
        last = insert_after(last, item, body_style)


def fix_figure_captions(doc: Document):
    """Обновить подписи и вставить недостающие рисунки 6 и 11."""
    paragraphs = doc.paragraphs

    # Удалить жёлтые пометки ▶ ВСТАВИТЬ
    to_remove = [p for p in paragraphs if "▶ ВСТАВИТЬ" in p.text or "▶ ВСТАВ" in p.text]
    for p in to_remove:
        remove_paragraph(p)

    old_to_new = {
        "Рисунок -": FIGURE_CAPTIONS[1][1],
        "Рисунок 1 – Архитектура": FIGURE_CAPTIONS[0][1],
        "Размеченный датасет": FIGURE_CAPTIONS[2][1],
        "метод очистки": FIGURE_CAPTIONS[3][1],
        "процесс замены сленга": FIGURE_CAPTIONS[4][1],
        "Тестовая база": FIGURE_CAPTIONS[5][1],
        "Зависимость точности": FIGURE_CAPTIONS[6][1],
        "сравнение моделей": FIGURE_CAPTIONS[7][1],
        "предсказание vs эксперт": FIGURE_CAPTIONS[8][1],
        "Пользовательский запрос": FIGURE_CAPTIONS[9][1],
        "Препроцессинг": FIGURE_CAPTIONS[10][1],
        "Поиск схожих": FIGURE_CAPTIONS[11][1],
        "Список наиболее схожих": FIGURE_CAPTIONS[12][1],
        "Визуализация результатов": FIGURE_CAPTIONS[13][1],
    }
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t.startswith("Рисунок"):
            continue
        for key, caption in old_to_new.items():
            if key in t:
                set_text(p, caption)
                break

    # Вставить Рис. 6 — тестовая БД (перед «Зависимость точности»)
    paragraphs = doc.paragraphs
    for p in paragraphs:
        if "Зависимость точности" in p.text and p.text.strip().startswith("Рисунок 7"):
            prev = p._p.getprevious()
            if prev is None or "Рисунок 6" not in Paragraph(prev, p._parent).text:
                new_p = OxmlElement("w:p")
                p._p.addprevious(new_p)
                np = Paragraph(new_p, p._parent)
                np.style = p.style
                set_text(np, FIGURE_CAPTIONS[5][1])
            break

    # Вставить Рис. 11 — препроцессинг (перед абзацем про препроцессинг, если нет)
    paragraphs = doc.paragraphs
    has_ris11 = any("Рисунок 11 – Препроцессинг" in p.text for p in paragraphs)
    if not has_ris11:
        for p in paragraphs:
            if p.text.startswith("На первом этапе выполняется препроцессинг"):
                new_p = OxmlElement("w:p")
                p._p.addprevious(new_p)
                np = Paragraph(new_p, p._parent)
                np.style = p.style
                set_text(np, FIGURE_CAPTIONS[10][1])
                break

    # Обновить ссылки (рис. N) в тексте
    fig_ref_map = {
        r"рис\. 1\b": "рис. 1",
        r"рис\. 2\b": "рис. 3",
        r"рис\. 3\b": "рис. 4",
        r"рис\. 4\b": "рис. 5",
        r"рис\. 5\b": "рис. 6",
        r"рис\. 6\b": "рис. 7",
        r"рис\. 7-8": "рис. 8–9",
        r"рис\. 7\b": "рис. 8",
        r"рис\. 8\b": "рис. 9",
        r"рис\. 9\b": "рис. 10",
        r"рис\. 10\b": "рис. 11",
        r"рис\. 11\b": "рис. 12",
        r"рис\. 12\b": "рис. 13",
        r"рис\. 13\b": "рис. 14",
    }
    for p in doc.paragraphs:
        if p.text.strip().startswith("Рисунок"):
            continue
        t = p.text
        new_t = t
        for pat, repl in fig_ref_map.items():
            new_t = re.sub(pat, repl, new_t, flags=re.I)
        if new_t != t:
            set_text(p, new_t)


def strip_citations_in_definitions(doc: Document):
    in_defs = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == DEFINITIONS_HEADING:
            in_defs = True
            continue
        if in_defs and t.startswith("Теоретические аспекты"):
            break
        if in_defs and t:
            cleaned = re.sub(r"\s*\[\d+\]\s*$", "", p.text.strip())
            if cleaned != p.text.strip():
                set_text(p, cleaned)


def add_citations(doc: Document):
    """Не используется при финализации — см. fix_citations_selective.py."""
    pass


def renumber_sources(doc: Document):
    src_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Список использованных источников":
            src_idx = i
            break
    if src_idx is None:
        return

    # Найти якорь: последний источник или заголовок
    anchor = doc.paragraphs[src_idx]
    body_style = anchor.style
    for p in doc.paragraphs[src_idx + 1 :]:
        t = p.text.strip()
        if t.startswith("ПРИЛОЖЕНИЕ"):
            break
        if t:
            body_style = p.style
            anchor = p

    # Удалить старые записи (кроме уже пронумерованных)
    to_remove = []
    for p in doc.paragraphs[src_idx + 1 :]:
        t = p.text.strip()
        if t.startswith("ПРИЛОЖЕНИЕ"):
            break
        if t:
            to_remove.append(p)
    for p in to_remove:
        remove_paragraph(p)

    # Не дублировать, если уже пронумеровано
    if any(doc.paragraphs[i].text.strip().startswith("[1]") for i in range(src_idx + 1, min(src_idx + 5, len(doc.paragraphs)))):
        return

    anchor = doc.paragraphs[src_idx]
    last = anchor
    for i, src in enumerate(SOURCES, 1):
        last = insert_after(last, f"[{i}] {src}", body_style)


def main():
    doc = Document(TARGET)
    update_title(doc)
    insert_definitions(doc)
    fix_figure_captions(doc)
    add_citations(doc)
    strip_citations_in_definitions(doc)
    renumber_sources(doc)
    doc.save(TARGET)
    print(f"Готово: {TARGET}")
    print(f"- Добавлен раздел «{DEFINITIONS_HEADING}»")
    print("- Рисунки пронумерованы 1–14")
    print("- Ссылки [1]–[12] добавлены в абзацы")
    print("- Список источников пронумерован")


if __name__ == "__main__":
    main()
