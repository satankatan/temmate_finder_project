# -*- coding: utf-8 -*-
"""Пометки для скриншотов, буквенные приложения А–Д, ссылки в тексте."""

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

TARGET = r"c:\Users\User\Downloads\teammate_finder_cherkashina_ktbo3_3.docx"

APPENDIX_MAP = {
    "1": "А",
    "2": "Б",
    "3": "В",
    "4": "Г",
    "5": "Д",
}

# Перед каждой подписью «Рисунок N» вставляется инструкция
FIGURE_MARKERS = {
    "Рисунок 1 – Архитектура модели Transformer": (
        "▶ ВСТАВИТЬ ИЗОБРАЖЕНИЕ (Рис. 1): Схема архитектуры Transformer — блоки Encoder/Decoder, "
        "механизм self-attention, поток данных. Можно взять из учебных материалов или нарисовать "
        "в draw.io / PowerPoint. Это теоретическая иллюстрация к §1.3."
    ),
    "Рисунок 2 – Размеченный датасет игровых описаний (gaming_compatibility_dataset.csv)": (
        "▶ ВСТАВИТЬ СКРИНШОТ (Рис. 2): Файл data/gaming_compatibility_dataset.csv открытый в Excel "
        "или VS Code. На экране должны быть видны столбцы: description1, description2, "
        "compatibility_score, game_type и 8–12 строк с примерами (в т.ч. со сленгом)."
    ),
    "Рисунок 3 – метод очистки предложений": (
        "▶ ВСТАВИТЬ СКРИНШОТ (Рис. 3): Фрагмент кода preprocessing/text_processor.py — "
        "метод clean_text() и этапы обработки (5 шагов). Снимок из VS Code / Cursor с подсветкой синтаксиса."
    ),
    "Рисунок 4 – процесс замены сленга": (
        "▶ ВСТАВИТЬ ИЗОБРАЖЕНИЕ (Рис. 4): Таблица «до / после» препроцессинга ИЛИ скрин словаря "
        "gaming_slang_dict в text_processor.py. Пример строк таблицы: "
        "«агр сап, давлю лайн» → «агрессивный саппорт давление линия»; "
        "«каррю на лейте» → «керри поздний фарм»."
    ),
    "Рисунок 5 – Тестовая база данных": (
        "▶ ВСТАВИТЬ СКРИНШОТ (Рис. 5): База teammate_finder.db в DB Browser for SQLite — "
        "таблица users с полями user_id, username, description, game_type (5–10 записей). "
        "Либо вывод команды после python scripts/seed_database.py."
    ),
    "Рисунок 6 – Зависимость точности от порога схожести": (
        "▶ ВСТАВИТЬ ИЗОБРАЖЕНИЕ (Рис. 6): График «Precision vs порог схожести». "
        "Файл из проекта: precision_vs_threshold.png (если есть) или построить в Excel "
        "по результатам тестов. Ось X — threshold (0.1–0.9), ось Y — precision."
    ),
    "Рисунок 7 – сравнение моделей языковой обработки по основным методам": (
        "▶ ВСТАВИТЬ ИЗОБРАЖЕНИЕ (Рис. 7): Столбчатая диаграмма сравнения TF-IDF / Sentence-BERT / "
        "Multilingual-BERT (Precision и NDCG). Файл: methods_comparison.png из папки проекта "
        "или скрин из sravnenie.py после запуска."
    ),
    "Рисунок 8 – корреляция моделей языковой обработки": (
        "▶ ВСТАВИТЬ ИЗОБРАЖЕНИЕ (Рис. 8): График корреляции предсказаний моделей с экспертными "
        "оценками. Можно использовать tsne_comparison.png или scatter-plot «predicted vs actual» "
        "из раздела сравнения методов."
    ),
    "Рисунок 9 – Пользовательский запрос": (
        "▶ ВСТАВИТЬ СКРИНШОТ (Рис. 9): Swagger UI (http://localhost:8000/docs) — эндпоинт "
        "POST /similarity/search с заполненным JSON (description, top_k, game_type) "
        "ИЛИ правая колонка teammate_finder.html с заполненным полем поиска."
    ),
    "Рисунок 10 – Препроцессинг входного текста": (
        "▶ ВСТАВИТЬ СКРИНШОТ (Рис. 10): Пример работы препроцессора — исходный текст со сленгом "
        "и результат clean_text (можно скрин консоли Python или таблица в Word). "
        "Пример: «агр сап, роамлю» → нормализованный текст."
    ),
    "Рисунок 11 – Поиск схожих пользователей на стороне бд": (
        "▶ ВСТАВИТЬ СКРИНШОТ (Рис. 11): Фрагмент database/sqlite_db.py — метод find_similar_users "
        "с вычислением косинусной схожести ИЛИ лог API с similarity scores при поиске."
    ),
    "Рисунок 12 – Список наиболее схожих пользователей": (
        "▶ ВСТАВИТЬ СКРИНШОТ (Рис. 12): JSON-ответ API /similarity/search в Swagger "
        "(массив user_id, username, description, similarity_score) после нажатия Execute."
    ),
    "Рисунок 13 – Визуализация результатов": (
        "▶ ВСТАВИТЬ СКРИНШОТ (Рис. 13): Главный экран teammate_finder.html — карточки найденных "
        "тиммейтов с полосой «Совместимость: XX%», панель статуса модели (загружена, словарь 240, "
        "размерность 128). Сделай поиск «Ищу агрессивного саппорта» после seed_database.py."
    ),
}

# Доп. пометка у описания модели (без отдельного номера в оглавлении)
MODEL_ARCH_MARKER = (
    "▶ ВСТАВИТЬ СХЕМУ (рядом с §2.2.3, по желанию как отдельный рисунок): "
    "Блок-схема TextEmbeddingModel: Текст → Препроцессор → Токены → Embedding → BiLSTM → "
    "Attention → FC → Вектор 128D. Можно нарисовать в draw.io."
)

TRAINING_MARKER = (
    "▶ ВСТАВИТЬ СКРИНШОТ (по желанию, §3.2): Консоль обучения python -m model.train — "
    "строки Epoch 20/20, train/val loss, «New best model saved», val_loss ≈ 0.1614."
)


def insert_paragraph_before(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    if text:
        run = new_para.add_run(text)
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        run.bold = True
    return new_para


def remove_old_appendix_hyperlinks(paragraph):
    """Удаляет гиперссылки «приложение 1»…«приложение 5» (остатки числовой нумерации)."""
    import re

    from docx.oxml.ns import qn

    for hyperlink in list(paragraph._p.findall(qn("w:hyperlink"))):
        link_text = "".join(node.text or "" for node in hyperlink.iter(qn("w:t")))
        if re.fullmatch(r"приложение [1-5]", link_text.strip(), re.I):
            paragraph._p.remove(hyperlink)


def set_paragraph_text(paragraph, text, highlight=False):
    remove_old_appendix_hyperlinks(paragraph)
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        run.bold = True


def fix_appendix_references(text):
    """Числовые приложения → буквенные."""
    import re

    text = re.sub(r"\(приложение ([АБВГД])\)приложение \d", r"(приложение \1)", text)
    text = text.replace("(приложение А)приложение 1", "(приложение А)")
    text = text.replace("(приложение Б)приложение 2", "(приложение Б)")
    text = text.replace("(приложение В)приложение 3", "(приложение В)")
    text = text.replace("(приложение Г)приложение 4", "(приложение Г)")

    fixes = [
        ("(приложение 1)", "(приложение А)"),
        ("(приложение 2)", "(приложение Б)"),
        ("(приложение 3)", "(приложение В)"),
        ("(приложение 4)", "(приложение Г)"),
        ("(приложение 5)", "(приложение Д)"),
        ("в приложении 5", "в приложении Д"),
        ("в приложении 4", "в приложении Г"),
        ("в приложении 3", "в приложении В"),
        ("в приложении 2", "в приложении Б"),
        ("в приложении 1", "в приложении А"),
        ("приложение 5", "приложение Д"),
        ("приложение 4", "приложение Г"),
        ("приложение 3", "приложение В"),
        ("приложение 2", "приложение Б"),
        ("приложение 1", "приложение А"),
        ("ПРИЛОЖЕНИЕ 5.", "ПРИЛОЖЕНИЕ Д."),
        ("ПРИЛОЖЕНИЕ 4.", "ПРИЛОЖЕНИЕ Г."),
        ("ПРИЛОЖЕНИЕ 3.", "ПРИЛОЖЕНИЕ В."),
        ("ПРИЛОЖЕНИЕ 2.", "ПРИЛОЖЕНИЕ Б."),
        ("ПРИЛОЖЕНИЕ 1.", "ПРИЛОЖЕНИЕ А."),
    ]
    for old, new in fixes:
        text = text.replace(old, new)
    return text


def main():
    doc = Document(TARGET)
    body_style = doc.paragraphs[84].style if len(doc.paragraphs) > 84 else None

    # 1. Ссылки на приложения в тексте (+ удаление старых гиперссылок)
    for p in doc.paragraphs:
        remove_old_appendix_hyperlinks(p)
        t = p.text
        if "приложен" in t.lower() and "ПРИЛОЖЕНИЕ" not in t and "приложений для поиска" not in t:
            new_t = fix_appendix_references(t)
            if new_t != t:
                set_paragraph_text(p, new_t)

    # 2. Заголовки приложений
    appendix_titles = {
        "ПРИЛОЖЕНИЕ 1. sqlite_db.py": "ПРИЛОЖЕНИЕ А. sqlite_db.py",
        "ПРИЛОЖЕНИЕ 1. Sqlite_db.py": "ПРИЛОЖЕНИЕ А. sqlite_db.py",
        "ПРИЛОЖЕНИЕ 2. text_processor.py": "ПРИЛОЖЕНИЕ Б. text_processor.py",
        "ПРИЛОЖЕНИЕ 3. network.py": "ПРИЛОЖЕНИЕ В. network.py",
        "ПРИЛОЖЕНИЕ 4. inference.py": "ПРИЛОЖЕНИЕ Г. inference.py",
        "ПРИЛОЖЕНИЕ 5. train.py": "ПРИЛОЖЕНИЕ Д. train.py",
    }
    for p in doc.paragraphs:
        t = p.text.strip()
        for old, new in appendix_titles.items():
            if old in t or t.startswith(old.split(".")[0]):
                if "sqlite" in t.lower() or "Sqlite" in t:
                    set_paragraph_text(p, "ПРИЛОЖЕНИЕ А. sqlite_db.py")
                elif "text_processor" in t:
                    set_paragraph_text(p, "ПРИЛОЖЕНИЕ Б. text_processor.py")
                elif "network" in t:
                    set_paragraph_text(p, "ПРИЛОЖЕНИЕ В. network.py")
                elif "inference" in t:
                    set_paragraph_text(p, "ПРИЛОЖЕНИЕ Г. inference.py")
                elif "train" in t:
                    set_paragraph_text(p, "ПРИЛОЖЕНИЕ Д. train.py")
                break

    # Исправить склеенный заголовок приложения А
    for p in doc.paragraphs:
        if "sqlite_db.pyПРИЛОЖЕНИЕ" in p.text or "sqlite_db.pyПРИЛОЖЕНИЕ" in p.text.replace(" ", ""):
            set_paragraph_text(p, "ПРИЛОЖЕНИЕ А. sqlite_db.py")

    # 3. Пометки для скриншотов (вставить перед подписью, если ещё нет)
    paragraphs = doc.paragraphs
    for p in paragraphs:
        t = p.text.strip()
        for caption, marker in FIGURE_MARKERS.items():
            if t == caption or t.startswith(caption.split("(")[0].strip()):
                prev_idx = None
                for i, para in enumerate(doc.paragraphs):
                    if para._p is p._p:
                        prev_idx = i - 1
                        break
                already = prev_idx is not None and prev_idx >= 0 and "▶ ВСТАВИТЬ" in doc.paragraphs[prev_idx].text
                if not already:
                    insert_paragraph_before(p, marker, body_style)
                break

    # 4. Доп. пометки у модели и обучения
    for p in doc.paragraphs:
        if p.text.startswith("Для задачи матчмейкинга разработана модель TextEmbeddingModel"):
            prev = p._p.getprevious()
            if prev is None or "▶ ВСТАВИТЬ СХЕМУ" not in (Paragraph(prev, p._parent).text if prev is not None else ""):
                insert_paragraph_before(p, MODEL_ARCH_MARKER, body_style)
        if "val_loss = 0.1614" in p.text or "val_loss ≈ 0.1614" in p.text:
            nxt = p._p.getnext()
            if nxt is None or "▶ ВСТАВИТЬ СКРИНШОТ" not in Paragraph(nxt, p._parent).text:
                insert_paragraph_after_simple(p, TRAINING_MARKER, body_style)

    doc.save(TARGET)
    print(f"Готово: {TARGET}")


def insert_paragraph_after_simple(paragraph, text, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    run = new_para.add_run(text)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.bold = True
    return new_para


if __name__ == "__main__":
    main()
