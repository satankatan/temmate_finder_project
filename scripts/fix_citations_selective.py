# -*- coding: utf-8 -*-
"""
Убрать ссылки [N] с каждого абзаца и оставить только там,
где текст реально опирается на конкретный источник.
"""

from __future__ import annotations

import re

from docx import Document

TARGET = r"c:\Users\User\Downloads\teammate_finder_cherkashina_ktbo3_3.docx"

DEFINITIONS_HEADING = "Определения, обозначения и сокращения"

# Уникальный фрагмент абзаца -> номер источника (только обоснованные ссылки)
SELECTIVE_CITATIONS = [
    # Введение
    ("Актуальность этой темы невозможно переоценить", "[1]"),
    ("Целью данной работы стала разработка интеллектуальной системы", "[8]"),
    ("Научная новизна заключается в разработке и обучении собственной", "[7]"),
    ("Разработанная система может быть внедрена в популярные игровые платформы", "[10]"),
    # Теория — психология
    ("Психологические критерии совместимости", None),  # заголовок — без ссылки
    ("противоречивых природных склонностей", "[11]"),
    ("психологической синергии", "[11]"),
    ("игровой психологии и современные технологии анализа данных", "[11]"),
    # Платформы
    ("Современные платформы для поиска тиммейтов", None),
    ("Большинство систем построено по принципу механического подбора", "[1]"),
    ("не по смысловому соответствию игровых философий", "[10]"),
    # NLP / Transformer
    ("Технологии обработки естественного языка для matching-систем", None),
    ("Основу современных matching-систем составляют методы преобразования", "[12]"),
    ("Традиционные методы TF-IDF", "[12]"),
    ("Модель машинного обучения Transformer", "[2]"),
    ("Для обучения модели сравнения текстов в проекте применяется Triplet Loss", "[7]"),
    # Практика — БД, препроцессор, модель
    ("Проектирование базы данных стало критически важным этапом", "[9]"),
    ("Традиционные методы очистки текста оказались неэффективны", "[5]"),
    ("Для задачи матчмейкинга разработана модель TextEmbeddingModel", "[12]"),
    ("Обучение выполнялось методом Triplet Loss", "[7]"),
    ("Основу алгоритма составляет вычисление косинусной схожести", "[8]"),
    # Тестирование и оценка
    ("Оценка эффективности проводилась в три этапа", "[8]"),
    ("Сравнительный анализ готовых моделей показал Precision 96.67% у Sentence-BERT", "[12]"),
    ("На отложенной тестовой выборке (11 пар, test.csv)", "[8]"),
    ("Сравнительный анализ baseline-методов выявил ограничения TF-IDF", "[12]"),
    ("Статистический анализ не выявил значимых различий между baseline", "[8]"),
    # Результат работы
    ("Алгоритм работы matching-системы реализует многоуровневый процесс", "[8]"),
    # Заключение
    ("Проделанная работа демонстрирует принципиальную возможность", "[12]"),
    ("Практическая значимость подтверждается работоспособностью", "[7]"),
    ("Перспективы развития: расширение размеченного датасета", "[7]"),
]


def set_text(paragraph, text: str):
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text


def strip_citation(text: str) -> str:
    return re.sub(r"\s*\[\d+\]\s*$", "", text.strip()).strip()


def main():
    doc = Document(TARGET)

    start = end = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Введение":
            start = i
        if p.text.strip() == "Список использованных источников":
            end = i
            break

    if start is None or end is None:
        raise RuntimeError("Не найдены границы основного текста")

    in_definitions = False
    removed = 0
    added = 0

    # 1. Убрать все ссылки из основного текста
    for p in doc.paragraphs[start:end]:
        t = p.text.strip()
        if t == DEFINITIONS_HEADING:
            in_definitions = True
            continue
        if in_definitions:
            if t.startswith("Теоретические аспекты"):
                in_definitions = False
            else:
                set_text(p, strip_citation(p.text))
                continue
        if not t or t.startswith("Рисунок") or "Heading" in p.style.name:
            continue
        if re.search(r"\[\d+\]\s*$", t):
            set_text(p, strip_citation(p.text))
            removed += 1

    # 2. Добавить только обоснованные ссылки
    for p in doc.paragraphs[start:end]:
        raw = strip_citation(p.text)
        if not raw or "Heading" in p.style.name:
            continue
        for fragment, cite in SELECTIVE_CITATIONS:
            if cite is None:
                continue
            if fragment in raw:
                if not raw.endswith(cite):
                    set_text(p, f"{raw} {cite}")
                    added += 1
                break

    doc.save(TARGET)
    print(f"Готово: {TARGET}")
    print(f"Удалено лишних ссылок: {removed}")
    print(f"Оставлено обоснованных: {added}")


if __name__ == "__main__":
    main()
